import io
import csv
import tempfile
from pathlib import Path
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.management import call_command
from django.test import TestCase, override_settings
from django.urls import reverse

from .models import (
    InterventionType,
    Project,
    ProjectMembership,
    PhotoDocumentation,
    RuianCadastralArea,
    RuianCadastralAreaMunicipality,
    RuianMunicipality,
    TreeAssessment,
    TreeIntervention,
    WorkRecord,
)


class ChangeConiferInterventionCommandTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(name="Conifer project")
        self.other_project = Project.objects.create(name="Other project")
        self.from_type, _ = InterventionType.objects.update_or_create(
            code="S-RZ",
            defaults={
                "name": "Řez zdravotní",
                "category": "Řez stromů",
            },
        )
        self.to_type, _ = InterventionType.objects.update_or_create(
            code="S-RB",
            defaults={
                "name": "Řez bezpečnostní",
                "category": "Řez stromů",
            },
        )

        self.conifer = WorkRecord.objects.create(
            project=self.project,
            passport_code="T-0001",
            taxon_latin="Picea abies",
            taxon="Picea abies",
            taxon_czech="smrk ztepilý",
        )
        self.czech_conifer = WorkRecord.objects.create(
            project=self.project,
            external_tree_id="EXT-2",
            taxon_czech="borovice lesní",
        )
        self.leaf = WorkRecord.objects.create(
            project=self.project,
            passport_code="T-0003",
            taxon_latin="Tilia cordata",
            taxon="Tilia cordata",
            taxon_czech="lípa srdčitá",
        )
        self.other_project_conifer = WorkRecord.objects.create(
            project=self.other_project,
            passport_code="T-0004",
            taxon_latin="Pinus sylvestris",
        )
        self.completed_conifer = WorkRecord.objects.create(
            project=self.project,
            passport_code="T-0005",
            taxon_latin="Abies alba",
        )

        self.conifer_intervention = TreeIntervention.objects.create(
            tree=self.conifer,
            intervention_type=self.from_type,
            status="proposed",
            urgency=1,
        )
        self.czech_conifer_intervention = TreeIntervention.objects.create(
            tree=self.czech_conifer,
            intervention_type=self.from_type,
            status="proposed",
            urgency=1,
        )
        self.leaf_intervention = TreeIntervention.objects.create(
            tree=self.leaf,
            intervention_type=self.from_type,
            status="proposed",
            urgency=1,
        )
        self.other_project_intervention = TreeIntervention.objects.create(
            tree=self.other_project_conifer,
            intervention_type=self.from_type,
            status="proposed",
            urgency=1,
        )
        self.completed_intervention = TreeIntervention.objects.create(
            tree=self.completed_conifer,
            intervention_type=self.from_type,
            status="completed",
            urgency=1,
        )

    def call_command(self, *extra_args):
        stdout = io.StringIO()
        call_command(
            "change_conifer_intervention",
            "--project-id",
            str(self.project.pk),
            "--from-code",
            "S-RZ",
            "--to-code",
            "S-RB",
            *extra_args,
            stdout=stdout,
        )
        return stdout.getvalue()

    def assert_type(self, intervention, intervention_type):
        intervention.refresh_from_db()
        self.assertEqual(intervention.intervention_type_id, intervention_type.pk)

    def test_dry_run_changes_nothing(self):
        output = self.call_command("--dry-run")

        self.assertIn("Dry run only", output)
        for intervention in (
            self.conifer_intervention,
            self.czech_conifer_intervention,
            self.leaf_intervention,
            self.other_project_intervention,
            self.completed_intervention,
        ):
            self.assert_type(intervention, self.from_type)

    def test_without_confirm_changes_nothing(self):
        self.call_command()

        self.assert_type(self.conifer_intervention, self.from_type)
        self.assert_type(self.czech_conifer_intervention, self.from_type)

    def test_confirm_changes_only_project_proposed_conifers(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            self.call_command("--backup-dir", tmpdir, "--confirm")

        self.assert_type(self.conifer_intervention, self.to_type)
        self.assert_type(self.czech_conifer_intervention, self.to_type)
        self.assert_type(self.leaf_intervention, self.from_type)
        self.assert_type(self.other_project_intervention, self.from_type)
        self.assert_type(self.completed_intervention, self.from_type)

    def test_include_completed_changes_completed_conifers(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            self.call_command("--backup-dir", tmpdir, "--confirm", "--include-completed")

        self.assert_type(self.conifer_intervention, self.to_type)
        self.assert_type(self.completed_intervention, self.to_type)

    def test_confirm_creates_csv_backup(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            self.call_command("--backup-dir", tmpdir, "--confirm")
            backup_files = list(Path(tmpdir).glob("change_conifer_intervention_project_*.csv"))
            self.assertEqual(len(backup_files), 1)

            with backup_files[0].open(encoding="utf-8-sig", newline="") as handle:
                rows = list(csv.DictReader(handle))

        self.assertEqual(len(rows), 2)
        intervention_ids = {int(row["intervention_id"]) for row in rows}
        self.assertEqual(
            intervention_ids,
            {self.conifer_intervention.pk, self.czech_conifer_intervention.pk},
        )
        for row in rows:
            self.assertEqual(row["project_id"], str(self.project.pk))
            self.assertEqual(row["old_code"], "S-RZ")
            self.assertEqual(row["new_code"], "S-RB")
            self.assertIn("estimated_price_czk", row)


class ExportProjectTreeCardsDocxCommandTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(name="DOCX projekt")
        self.other_project = Project.objects.create(name="Jiný projekt")
        self.intervention_type, _ = InterventionType.objects.update_or_create(
            code="S-RB",
            defaults={
                "name": "Řez bezpečnostní",
                "category": "Řez stromů",
            },
        )
        self.tree = WorkRecord.objects.create(
            project=self.project,
            passport_code="T-0002",
            taxon_latin="Picea abies",
            taxon_czech="smrk ztepilý",
        )
        self.tree_without_comment = WorkRecord.objects.create(
            project=self.project,
            passport_code="T-0003",
            taxon="Tilia cordata",
        )
        self.other_tree = WorkRecord.objects.create(
            project=self.other_project,
            passport_code="T-9999",
            taxon="Pinus sylvestris",
        )
        self.project.trees.add(self.tree, self.tree_without_comment)
        self.other_project.trees.add(self.other_tree)

        TreeIntervention.objects.create(
            tree=self.tree,
            intervention_type=self.intervention_type,
            description="Odstranit suché větve.",
            status="proposed",
        )
        TreeIntervention.objects.create(
            tree=self.tree,
            intervention_type=self.intervention_type,
            description="Odstranit suché větve.",
            status="completed",
        )
        TreeIntervention.objects.create(
            tree=self.tree_without_comment,
            intervention_type=self.intervention_type,
            description="",
            status="proposed",
        )
        TreeIntervention.objects.create(
            tree=self.other_tree,
            intervention_type=self.intervention_type,
            description="Nemá být exportováno.",
            status="proposed",
        )

    def test_command_creates_docx_for_project_trees(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "cards.docx"
            call_command(
                "export_project_tree_cards_docx",
                "--project-id",
                str(self.project.pk),
                "--output",
                str(output_path),
            )

            self.assertTrue(output_path.exists())
            document = Document(output_path)

        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        full_text = "\n".join(paragraphs)
        self.assertIn("T-0002", full_text)
        self.assertIn("Taxon: smrk ztepilý (Picea abies)", full_text)
        self.assertIn("Komentář", full_text)
        self.assertIn("Řez bezpečnostní: Odstranit suché větve.", full_text)
        self.assertEqual(full_text.count("Odstranit suché větve."), 1)
        self.assertIn("T-0003", full_text)
        self.assertIn("Řez bezpečnostní", full_text)
        self.assertNotIn("Řez bezpečnostní:", "\n".join(
            paragraph.text
            for paragraph in document.paragraphs[
                paragraphs.index("T-0003"):
            ]
        ))
        self.assertIn("Fotografie není k dispozici", full_text)
        self.assertNotIn("Komentář není vyplněn", full_text)
        self.assertNotIn("T-9999", full_text)
        self.assertNotIn("Nemá být exportováno.", full_text)

    def test_batch_export_creates_numbered_docx_files(self):
        from docx import Document

        extra_tree = WorkRecord.objects.create(
            project=self.project,
            passport_code="T-0004",
            taxon="Acer platanoides",
        )
        self.project.trees.add(extra_tree)

        with tempfile.TemporaryDirectory() as tmpdir:
            stdout = io.StringIO()
            call_command(
                "export_project_tree_cards_docx",
                "--project-id",
                str(self.project.pk),
                "--output-dir",
                tmpdir,
                "--batch-size",
                "2",
                stdout=stdout,
            )

            files = sorted(path.name for path in Path(tmpdir).glob("*.docx"))
            self.assertEqual(
                files,
                [
                    f"project_{self.project.pk}_stromy_001_002.docx",
                    f"project_{self.project.pk}_stromy_003_003.docx",
                ],
            )
            first_document = Document(Path(tmpdir) / files[0])
            second_document = Document(Path(tmpdir) / files[1])

        self.assertIn("Exported 2/3 trees...", stdout.getvalue())
        self.assertIn("Exported 3/3 trees...", stdout.getvalue())
        self.assertIn(
            "T-0002",
            "\n".join(paragraph.text for paragraph in first_document.paragraphs),
        )
        self.assertIn(
            "T-0004",
            "\n".join(paragraph.text for paragraph in second_document.paragraphs),
        )

    def test_image_size_is_limited_by_width_and_height(self):
        from docx.shared import Cm
        from PIL import Image

        from tracker.management.commands.export_project_tree_cards_docx import (
            _batch_filename,
            _fit_image_size,
        )

        self.assertEqual(
            _batch_filename(37, 1, 100),
            "project_37_stromy_001_100.docx",
        )

        portrait = io.BytesIO()
        Image.new("RGB", (800, 2400), "white").save(portrait, format="JPEG")
        width, height = _fit_image_size(portrait.getvalue(), Cm(9), Cm(9))
        self.assertLessEqual(int(width), int(Cm(9)))
        self.assertLessEqual(int(height), int(Cm(9)))
        self.assertEqual(int(height), int(Cm(9)))

        landscape = io.BytesIO()
        Image.new("RGB", (2400, 800), "white").save(landscape, format="JPEG")
        width, height = _fit_image_size(landscape.getvalue(), Cm(9), Cm(9))
        self.assertLessEqual(int(width), int(Cm(9)))
        self.assertLessEqual(int(height), int(Cm(9)))
        self.assertEqual(int(width), int(Cm(9)))


class ProjectTreeAddTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="worker", password="pass1234"
        )
        self.foreman = get_user_model().objects.create_user(
            username="foreman", password="pass1234"
        )
        self.project = Project.objects.create(name="P1")
        ProjectMembership.objects.create(
            user=self.user, project=self.project, role=ProjectMembership.Role.WORKER
        )
        ProjectMembership.objects.create(
            user=self.foreman,
            project=self.project,
            role=ProjectMembership.Role.FOREMAN,
        )
        self.work_record = WorkRecord.objects.create(
            title="WR-1", latitude=49.0, longitude=17.0
        )

    def test_project_tree_add_requires_foreman(self):
        self.client.force_login(self.user)
        url = reverse("project_tree_add", args=[self.project.pk, self.work_record.pk])
        resp = self.client.post(url)
        self.assertEqual(resp.status_code, 403)

    def test_project_tree_add_sets_fk_if_null(self):
        self.client.force_login(self.foreman)
        url = reverse("project_tree_add", args=[self.project.pk, self.work_record.pk])
        resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        self.work_record.refresh_from_db()
        self.assertTrue(self.project.trees.filter(pk=self.work_record.pk).exists())
        self.assertEqual(self.work_record.project_id, self.project.pk)


@override_settings(STATICFILES_STORAGE="django.contrib.staticfiles.storage.StaticFilesStorage")
class ProjectTreeListSmokeTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="project-user", password="pass1234"
        )
        self.project = Project.objects.create(name="P1", description="Demo project")
        ProjectMembership.objects.create(
            user=self.user,
            project=self.project,
            role=ProjectMembership.Role.WORKER,
        )
        self.oak = WorkRecord.objects.create(title="Oak 001")
        self.pine = WorkRecord.objects.create(title="Pine 002")
        self.project.trees.add(self.oak, self.pine)

    def test_project_tree_list_returns_200(self):
        self.client.force_login(self.user)
        response = self.client.get(reverse("project_tree_list", args=[self.project.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Oak 001")
        self.assertContains(response, "Pine 002")

    def test_project_tree_list_filters_by_q(self):
        self.client.force_login(self.user)
        response = self.client.get(
            reverse("project_tree_list", args=[self.project.pk]),
            {"q": "Oak"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Oak 001")
        self.assertNotContains(response, "Pine 002")

    def test_project_tree_list_items_returns_partial(self):
        self.client.force_login(self.user)
        response = self.client.get(
            reverse("project_tree_list_items", args=[self.project.pk])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "workrecord-row")
        self.assertContains(response, "Oak 001")

    def test_project_detail_links_to_project_tree_list(self):
        self.client.force_login(self.user)
        response = self.client.get(reverse("project_detail", args=[self.project.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Seznam stromů")
        self.assertContains(response, reverse("project_tree_list", args=[self.project.pk]))


class ProjectXlsxExportTests(TestCase):
    def test_display_without_code_strips_only_matching_prefix(self):
        from tracker.views import _display_without_code

        self.assertEqual(
            _display_without_code(3, "3 – výrazně snížená"),
            "výrazně snížená",
        )
        self.assertEqual(
            _display_without_code("a", "a – dlouhodobě perspektivní"),
            "dlouhodobě perspektivní",
        )
        self.assertEqual(
            _display_without_code(1, "1 Pomístní překážky"),
            "Pomístní překážky",
        )
        self.assertEqual(
            _display_without_code(1, "Pomístní překážky"),
            "Pomístní překážky",
        )

    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="xlsx-user", password="pass1234"
        )
        self.outsider = get_user_model().objects.create_user(
            username="xlsx-outsider", password="pass1234"
        )
        self.project = Project.objects.create(name="Městský park")
        self.legacy_project = Project.objects.create(name="Legacy projekt")
        ProjectMembership.objects.create(
            user=self.user,
            project=self.project,
            role=ProjectMembership.Role.WORKER,
        )
        self.first_tree = WorkRecord.objects.create(
            title="T-001",
            description="Běžný popis",
            project=self.legacy_project,
            vegetation_type=WorkRecord.VegetationType.TREE,
        )
        self.second_tree = WorkRecord.objects.create(
            title="T-002",
            description="",
            project=None,
            vegetation_type=WorkRecord.VegetationType.TREE,
        )
        self.project.trees.add(self.first_tree, self.second_tree)
        TreeAssessment.objects.create(
            work_record=self.first_tree,
            vitality=2,
            stability=1,
            access_obstacle_level=1,
        )
        intervention_type = InterventionType.objects.create(
            code="RZ",
            name="Redukční řez",
            category="Řez",
            description="Redukce koruny stromu",
        )
        self.first_intervention = TreeIntervention.objects.create(
            tree=self.first_tree,
            intervention_type=intervention_type,
            description="Odstranit suché větve",
            urgency=1,
            status="proposed",
            status_note="Nutná kontrola",
            assigned_to=self.user,
        )
        TreeIntervention.objects.filter(pk=self.first_intervention.pk).update(
            estimated_price_czk=1500
        )
        self.first_intervention.refresh_from_db()
        urgent_type = InterventionType.objects.create(
            code="KONT",
            name="Bezpečnostní kontrola",
            category="Kontrola",
        )
        self.urgent_intervention = TreeIntervention.objects.create(
            tree=self.first_tree,
            intervention_type=urgent_type,
            description="Odstranit suché větve",
            urgency=0,
            status="proposed",
            status_note="Nutná kontrola",
        )
        TreeIntervention.objects.filter(pk=self.urgent_intervention.pk).update(
            estimated_price_czk=500
        )
        self.urgent_intervention.refresh_from_db()
        self.second_intervention = TreeIntervention.objects.create(
            tree=self.second_tree,
            intervention_type=intervention_type,
            description="Kontrolní řez",
            urgency=3,
            status="completed",
        )
        TreeIntervention.objects.filter(pk=self.second_intervention.pk).update(
            estimated_price_czk=800
        )
        self.second_intervention.refresh_from_db()
        PhotoDocumentation.objects.create(
            work_record=self.first_tree,
            photo="photos/first.jpg",
            description="Celek",
        )
        PhotoDocumentation.objects.create(
            work_record=self.second_tree,
            photo="photos/second.jpg",
            description="Detail",
        )
        self.url = reverse("export_selected_xlsx", args=[self.project.pk])

    def _workbook(self, response):
        from openpyxl import load_workbook

        return load_workbook(io.BytesIO(response.content), data_only=False)

    def _rows_by_header(self, worksheet, key_header):
        headers = [cell.value for cell in worksheet[1]]
        key_index = headers.index(key_header)
        return {
            row[key_index].value: {
                headers[index]: cell.value for index, cell in enumerate(row)
            }
            for row in worksheet.iter_rows(min_row=2)
        }

    def test_export_all_project_has_expected_sheets_and_overview(self):
        self.client.force_login(self.user)

        response = self.client.post(self.url, {"export_all": "1"})

        self.assertEqual(response.status_code, 200)
        self.assertIn(
            "arbomap_mestsky_park_",
            response["Content-Disposition"],
        )
        self.assertIn("_cely-projekt.xlsx", response["Content-Disposition"])
        workbook = self._workbook(response)
        self.assertEqual(
            workbook.sheetnames,
            ["Přehled stromů", "Zásahy", "Fotky"],
        )
        self.assertNotIn("Technická data", workbook.sheetnames)
        worksheet = workbook["Přehled stromů"]
        headers = [cell.value for cell in worksheet[1]]
        self.assertEqual(
            headers,
            [
                "Číslo stromu",
                "Taxon",
                "Český název",
                "Datum hodnocení",
                "Výška [m]",
                "Obvod kmene [cm]",
                "DBH [cm]",
                "Šířka koruny [m]",
                "Plocha koruny [m²]",
                "Fyziologické stáří",
                "Fyziologické stáří slovně",
                "Vitalita",
                "Vitalita slovně",
                "Zdravotní stav",
                "Zdravotní stav slovně",
                "Stabilita",
                "Stabilita slovně",
                "Překážka",
                "Překážka slovně",
                "Perspektiva",
                "Perspektiva slovně",
                "Jmelí",
                "Jmelí slovně",
                "Navržené zásahy",
                "Naléhavost zásahu",
                "Naléhavost zásahu slovně",
                "Odhadovaná cena zásahů",
                "Poznámka k zásahům",
                "GPS šířka",
                "GPS délka",
                "Parcela",
                "Katastrální území",
                "Obec",
            ],
        )
        self.assertNotIn("Název", headers)
        self.assertNotIn("Typ vegetace", headers)
        self.assertNotIn("Latinský název", headers)
        self.assertIn("Taxon", headers)
        self.assertIn("Naléhavost zásahu", headers)
        self.assertIn("Naléhavost zásahu slovně", headers)
        self.assertIn("Odhadovaná cena zásahů", headers)
        self.assertIn("Poznámka k zásahům", headers)
        self.assertIn("Překážka", headers)
        self.assertIn("Překážka slovně", headers)
        self.assertNotIn("Počet fotek", headers)
        self.assertNotIn("Stav zásahů", headers)
        self.assertEqual(
            headers[-5:],
            ["GPS šířka", "GPS délka", "Parcela", "Katastrální území", "Obec"],
        )
        rows = self._rows_by_header(worksheet, "Číslo stromu")
        self.assertEqual(set(rows), {"T-001", "T-002"})
        self.assertEqual(rows["T-001"]["Vitalita"], 2)
        self.assertEqual(rows["T-001"]["Vitalita slovně"], "zřetelně snížená")
        self.assertEqual(rows["T-001"]["Stabilita"], 1)
        self.assertEqual(
            rows["T-001"]["Stabilita slovně"],
            "výborná až dobrá (nenarušená)",
        )
        self.assertEqual(rows["T-001"]["Překážka"], 1)
        self.assertEqual(rows["T-001"]["Překážka slovně"], "Pomístní překážky")
        self.assertIn("RZ – Redukční řez", rows["T-001"]["Navržené zásahy"])
        self.assertEqual(rows["T-001"]["Naléhavost zásahu"], 0)
        self.assertEqual(
            rows["T-001"]["Naléhavost zásahu slovně"],
            "okamžitě, riziko z prodlení",
        )
        self.assertEqual(rows["T-001"]["Odhadovaná cena zásahů"], 2000)
        self.assertEqual(
            rows["T-001"]["Poznámka k zásahům"],
            "Odstranit suché větve; Nutná kontrola",
        )
        self.assertEqual(worksheet.freeze_panes, "A2")
        self.assertTrue(worksheet.auto_filter.ref)
        self.assertTrue(all(cell.font.bold for cell in worksheet[1]))
        self.assertTrue(
            all(
                cell.alignment.vertical == "top"
                for row in worksheet.iter_rows(min_row=2)
                for cell in row
            )
        )

    def test_intervention_and_photo_are_exported_to_detail_sheets(self):
        self.client.force_login(self.user)

        workbook = self._workbook(
            self.client.post(self.url, {"export_all": "1"})
        )
        intervention_rows = self._rows_by_header(
            workbook["Zásahy"],
            "Název stromu",
        )
        photo_rows = self._rows_by_header(
            workbook["Fotky"],
            "Název stromu",
        )

        self.assertEqual(
            intervention_rows["T-001"]["Kód zásahu"],
            "RZ",
        )
        self.assertEqual(
            intervention_rows["T-001"]["Odhad ceny"],
            1500,
        )
        self.assertEqual(
            intervention_rows["T-001"]["Zodpovědná osoba"],
            self.user.get_username(),
        )
        intervention_headers = [
            cell.value for cell in workbook["Zásahy"][1]
        ]
        self.assertIn("Konkrétní popis", intervention_headers)
        self.assertIn("Stavová poznámka", intervention_headers)
        self.assertIn("Stav", intervention_headers)
        self.assertEqual(photo_rows["T-001"]["Název souboru"], "first.jpg")
        photo_sheet = workbook["Fotky"]
        photo_headers = [cell.value for cell in photo_sheet[1]]
        tree_name_column = photo_headers.index("Název stromu") + 1
        first_tree_row = next(
            row_index
            for row_index in range(2, photo_sheet.max_row + 1)
            if photo_sheet.cell(row=row_index, column=tree_name_column).value == "T-001"
        )
        link_cell = photo_sheet.cell(
            row=first_tree_row,
            column=photo_headers.index("URL fotky") + 1,
        )
        self.assertEqual(link_cell.value, "Otevřít fotku")
        self.assertTrue(link_cell.hyperlink)
        self.assertIn("photos/first.jpg", link_cell.hyperlink.target)

    def test_export_selected_contains_only_selected_tree_on_all_sheets(self):
        self.client.force_login(self.user)

        response = self.client.post(
            self.url,
            {"export_selected": "1", "selected_records": [self.first_tree.pk]},
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn("_vyber.xlsx", response["Content-Disposition"])
        workbook = self._workbook(response)
        sheet_keys = {
            "Přehled stromů": "Číslo stromu",
            "Zásahy": "Název stromu",
            "Fotky": "Název stromu",
        }
        for sheet_name, key_header in sheet_keys.items():
            rows = self._rows_by_header(workbook[sheet_name], key_header)
            self.assertEqual(set(rows), {"T-001"}, sheet_name)

    def test_user_without_project_access_does_not_receive_workbook(self):
        self.client.force_login(self.outsider)

        response = self.client.post(self.url, {"export_all": "1"})

        self.assertEqual(response.status_code, 302)
        self.assertNotEqual(
            response.get("Content-Type"),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    def test_empty_values_export_without_error(self):
        self.client.force_login(self.user)

        response = self.client.post(
            self.url,
            {"export_selected": "1", "selected_records": [self.second_tree.pk]},
        )

        self.assertEqual(response.status_code, 200)
        workbook = self._workbook(response)
        overview_row = self._rows_by_header(
            workbook["Přehled stromů"],
            "Číslo stromu",
        )["T-002"]
        self.assertIsNone(overview_row["Datum hodnocení"])

    def test_formula_like_user_text_is_exported_as_text(self):
        self.first_tree.taxon = "=2+2"
        self.first_tree.save(update_fields=["taxon"])
        self.client.force_login(self.user)

        response = self.client.post(
            self.url,
            {"export_selected": "1", "selected_records": [self.first_tree.pk]},
        )

        worksheet = self._workbook(response)["Přehled stromů"]
        headers = [cell.value for cell in worksheet[1]]
        taxon_cell = worksheet.cell(
            row=2,
            column=headers.index("Taxon") + 1,
        )
        self.assertEqual(taxon_cell.value, "'=2+2")
        self.assertEqual(taxon_cell.data_type, "s")


class WorkrecordsGeojsonTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="user1", password="pass1234"
        )
        self.project1 = Project.objects.create(name="P1")
        self.project2 = Project.objects.create(name="P2")
        ProjectMembership.objects.create(
            user=self.user,
            project=self.project1,
            role=ProjectMembership.Role.WORKER,
        )
        self.in_project = WorkRecord.objects.create(
            title="InProject", latitude=49.1, longitude=17.1
        )
        self.in_project_far = WorkRecord.objects.create(
            title="InProjectFar", latitude=49.8, longitude=17.8, vegetation_type=None
        )
        self.in_project_blank = WorkRecord.objects.create(
            title="InProjectBlank", latitude=49.7, longitude=17.7, vegetation_type=""
        )
        self.other_project = WorkRecord.objects.create(
            title="OtherProject", latitude=49.2, longitude=17.2
        )
        self.orphan = WorkRecord.objects.create(
            title="Orphan", latitude=49.3, longitude=17.3
        )
        self.project1.trees.add(self.in_project)
        self.project1.trees.add(self.in_project_far)
        self.project1.trees.add(self.in_project_blank)
        self.project2.trees.add(self.other_project)

    def test_workrecords_geojson_non_project_uses_m2m(self):
        self.client.force_login(self.user)
        url = reverse("workrecords_geojson")
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        payload = resp.json()
        ids = {f["properties"]["id"] for f in payload.get("features", [])}
        self.assertIn(self.in_project.pk, ids)
        self.assertIn(self.in_project_far.pk, ids)
        self.assertIn(self.in_project_blank.pk, ids)
        self.assertNotIn(self.other_project.pk, ids)
        self.assertNotIn(self.orphan.pk, ids)

    def test_workrecords_geojson_project_intervention_types_ignore_bbox(self):
        type_a = InterventionType.objects.create(code="A", name="Typ A")
        type_b = InterventionType.objects.create(code="B", name="Typ B")
        TreeIntervention.objects.create(tree=self.in_project, intervention_type=type_a)
        TreeIntervention.objects.create(
            tree=self.in_project_far,
            intervention_type=type_b,
            status="completed",
        )

        self.client.force_login(self.user)
        url = reverse("workrecords_geojson")
        resp = self.client.get(
            url,
            {
                "project": self.project1.pk,
                "bbox": "17.05,49.05,17.2,49.2",
            },
        )
        self.assertEqual(resp.status_code, 200)
        payload = resp.json()
        ids = {f["properties"]["id"] for f in payload.get("features", [])}
        self.assertEqual(ids, {self.in_project.pk})
        self.assertEqual(payload.get("project_intervention_types"), ["A", "B"])
        self.assertEqual(payload.get("project_intervention_type_counts"), {"A": 1, "B": 1})
        self.assertEqual(payload.get("project_intervention_statuses"), ["proposed", "completed"])
        self.assertEqual(
            payload.get("project_intervention_status_counts"),
            {"proposed": 1, "completed": 1},
        )
        self.assertEqual(payload.get("project_vegetation_counts", {}).get("TREE"), 3)
        self.assertEqual(payload.get("project_no_intervention_count"), 1)
        feature = payload["features"][0]
        self.assertEqual(feature["properties"]["intervention_types"], ["A"])
        self.assertEqual(feature["properties"]["intervention_statuses"], ["proposed"])
        self.assertTrue(feature["properties"]["has_interventions"])

    def test_workrecords_geojson_marks_only_non_completed_interventions_active(self):
        type_prune = InterventionType.objects.create(code="PR", name="Řez")
        type_removal = InterventionType.objects.create(
            code="KAC", name="Kácení", category="Kácení"
        )
        TreeIntervention.objects.create(
            tree=self.in_project,
            intervention_type=type_prune,
            status="completed",
        )
        TreeIntervention.objects.create(
            tree=self.in_project_far,
            intervention_type=type_prune,
            status="done_pending_owner",
        )
        TreeIntervention.objects.create(
            tree=self.in_project_blank,
            intervention_type=type_removal,
            status="proposed",
        )

        self.client.force_login(self.user)
        resp = self.client.get(reverse("workrecords_geojson"), {"project": self.project1.pk})
        self.assertEqual(resp.status_code, 200)
        features_by_id = {
            feature["properties"]["id"]: feature["properties"]
            for feature in resp.json().get("features", [])
        }

        self.assertFalse(features_by_id[self.in_project.pk]["has_active_intervention"])
        self.assertFalse(features_by_id[self.in_project.pk]["has_removal_intervention"])
        self.assertEqual(features_by_id[self.in_project.pk]["intervention_stage"], "approved")

        self.assertTrue(features_by_id[self.in_project_far.pk]["has_active_intervention"])
        self.assertFalse(features_by_id[self.in_project_far.pk]["has_removal_intervention"])
        self.assertEqual(features_by_id[self.in_project_far.pk]["intervention_stage"], "done")

        self.assertTrue(features_by_id[self.in_project_blank.pk]["has_active_intervention"])
        self.assertTrue(features_by_id[self.in_project_blank.pk]["has_removal_intervention"])
        self.assertEqual(features_by_id[self.in_project_blank.pk]["intervention_stage"], "none")


class CadastreAreaCodeDerivationTests(TestCase):
    def _create_with_parcel(self, parcel_number):
        with patch(
            "tracker.models._cad_lookup_by_point",
            return_value={"parcel_number": parcel_number, "cad_lookup_status": "ok"},
        ):
            record = WorkRecord.objects.create(
                title="WR", latitude=49.0, longitude=17.0
            )
        record.refresh_from_db()
        return record

    def test_cadastral_area_code_derived_from_parcel_number(self):
        record = self._create_with_parcel("710504-241/1")
        self.assertEqual(record.cadastral_area_code, "710504")
        record = self._create_with_parcel("713520-300/2")
        self.assertEqual(record.cadastral_area_code, "713520")

    def test_cadastral_area_code_not_set_for_invalid_parcel_number(self):
        record = self._create_with_parcel("")
        self.assertFalse(record.cadastral_area_code)
        record = self._create_with_parcel("invalid")
        self.assertFalse(record.cadastral_area_code)

    def test_cadastral_area_code_accepts_numeric_prefix(self):
        record = self._create_with_parcel("123-")
        self.assertEqual(record.cadastral_area_code, "123")


class RuianImportTests(TestCase):
    def test_import_ruian_from_sample_dir(self):
        sample_dir = Path(__file__).resolve().parent / "data" / "ruian_sample"
        call_command("import_ruian", source_dir=str(sample_dir))

        self.assertEqual(RuianMunicipality.objects.count(), 2)
        self.assertEqual(RuianCadastralArea.objects.count(), 3)
        self.assertTrue(
            RuianCadastralAreaMunicipality.objects.filter(
                cadastral_area_id="2001", municipality_id="1001"
            ).exists()
        )


class WorkRecordRuianEnrichmentTests(TestCase):
    def _create_with_parcel(self, parcel_number):
        with patch(
            "tracker.models._cad_lookup_by_point",
            return_value={"parcel_number": parcel_number, "cad_lookup_status": "ok"},
        ):
            record = WorkRecord.objects.create(
                title="WR", latitude=49.0, longitude=17.0
            )
        record.refresh_from_db()
        return record

    def test_enriches_names_from_ruian_tables(self):
        RuianCadastralArea.objects.create(code="710504", name="Katastr X")
        RuianMunicipality.objects.create(code="5001", name="Obec Y")
        RuianCadastralAreaMunicipality.objects.create(
            cadastral_area_id="710504", municipality_id="5001"
        )

        record = self._create_with_parcel("710504-241/1")
        self.assertEqual(record.cadastral_area_code, "710504")
        self.assertEqual(record.cadastral_area_name, "Katastr X")
        self.assertEqual(record.municipality_name, "Obec Y")

    def test_missing_ruian_tables_does_not_block_create(self):
        record = self._create_with_parcel("710504-241/1")
        self.assertEqual(record.cadastral_area_code, "710504")
        self.assertFalse(record.cadastral_area_name)
        self.assertFalse(record.municipality_name)
