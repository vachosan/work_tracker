import io
import re
from pathlib import Path

from django.core.management.base import BaseCommand, CommandError
from django.db.models import Prefetch

from tracker.models import PhotoDocumentation, Project, TreeIntervention


def _natural_sort_key(value):
    text = str(value or "")
    return [
        int(part) if part.isdigit() else part.lower()
        for part in re.split(r"(\d+)", text)
    ]


def _tree_sort_key(tree):
    label = tree.preferred_id_label
    return (
        _natural_sort_key(label),
        tree.pk or 0,
    )


def _taxon_label(tree):
    if tree.taxon_czech and tree.taxon_latin:
        return f"{tree.taxon_czech} ({tree.taxon_latin})"
    return tree.taxon_czech or tree.taxon_latin or tree.taxon or ""


def _intervention_label(intervention):
    intervention_type = intervention.intervention_type
    if not intervention_type:
        return ""
    return intervention_type.name or intervention_type.code or ""


def _comment_paragraphs(tree):
    comments = []
    seen = set()
    for intervention in getattr(tree, "export_interventions", []):
        comment = (intervention.description or "").strip()
        label = _intervention_label(intervention).strip()
        if not label and not comment:
            continue

        normalized_label = re.sub(r"\s+", " ", label).casefold()
        normalized_comment = re.sub(r"\s+", " ", comment).casefold()
        normalized = (normalized_label, normalized_comment)
        if normalized in seen:
            continue
        seen.add(normalized)

        if label and comment:
            comments.append(f"{label}: {comment}")
        elif label:
            comments.append(label)
        else:
            comments.append(comment)
    return comments


def _first_photo(tree):
    for photo in getattr(tree, "export_photos", []):
        if photo.photo:
            return photo
    return None


def _fit_image_size(image_bytes, max_width, max_height):
    from docx.shared import Emu
    from PIL import Image

    with Image.open(io.BytesIO(image_bytes)) as image:
        pixel_width, pixel_height = image.size

    if pixel_width <= 0 or pixel_height <= 0:
        return None, None

    scale = min(int(max_width) / pixel_width, int(max_height) / pixel_height)
    return (
        Emu(round(pixel_width * scale)),
        Emu(round(pixel_height * scale)),
    )


def _batch_filename(project_id, first_number, last_number):
    return f"project_{project_id}_stromy_{first_number:03d}_{last_number:03d}.docx"


class Command(BaseCommand):
    help = "Export all project trees into a DOCX with one tree card per page."

    def add_arguments(self, parser):
        parser.add_argument("--project-id", type=int, required=True)
        parser.add_argument("--output")
        parser.add_argument("--output-dir")
        parser.add_argument("--batch-size", type=int, default=100)

    def handle(self, *args, **options):
        try:
            from docx.shared import Inches
        except ModuleNotFoundError as exc:
            raise CommandError(
                "Missing dependency python-docx. Install requirements or run: pip install python-docx"
            ) from exc

        project_id = options["project_id"]
        output = options.get("output")
        output_dir = options.get("output_dir")
        batch_size = options["batch_size"]

        if bool(output) == bool(output_dir):
            raise CommandError("Pass exactly one of --output or --output-dir.")
        if batch_size <= 0:
            raise CommandError("--batch-size must be greater than zero.")

        try:
            project = Project.objects.get(pk=project_id)
        except Project.DoesNotExist as exc:
            raise CommandError(f"Project {project_id} does not exist.") from exc

        trees = list(
            project.trees.all()
            .prefetch_related(
                Prefetch(
                    "photos",
                    queryset=PhotoDocumentation.objects.order_by("id"),
                    to_attr="export_photos",
                ),
                Prefetch(
                    "interventions",
                    queryset=TreeIntervention.objects.select_related(
                        "intervention_type"
                    ).order_by("status", "urgency", "due_date", "id"),
                    to_attr="export_interventions",
                ),
            )
        )
        trees.sort(key=_tree_sort_key)

        if output:
            output_path = Path(output)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self._export_document(trees, output_path)
            self.stdout.write(
                self.style.SUCCESS(f"Exported {len(trees)} trees to {output_path}")
            )
            return

        output_dir_path = Path(output_dir)
        output_dir_path.mkdir(parents=True, exist_ok=True)
        total = len(trees)
        if not total:
            self.stdout.write(self.style.SUCCESS(f"Exported 0 trees to {output_dir_path}"))
            return
        for start in range(0, total, batch_size):
            batch = trees[start : start + batch_size]
            first_number = start + 1
            last_number = start + len(batch)
            output_path = output_dir_path / _batch_filename(
                project_id,
                first_number,
                last_number,
            )
            self._export_document(batch, output_path)
            self.stdout.write(f"Exported {last_number}/{total} trees...")

        self.stdout.write(
            self.style.SUCCESS(f"Exported {total} trees to {output_dir_path}")
        )

    def _export_document(self, trees, output_path):
        from docx import Document
        from docx.shared import Cm, Inches

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        for index, tree in enumerate(trees):
            if index:
                document.add_page_break()

            tree_label = tree.preferred_id_label
            document.add_heading(tree_label, level=1)

            taxon = _taxon_label(tree)
            if taxon:
                paragraph = document.add_paragraph()
                paragraph.add_run("Taxon: ").bold = True
                paragraph.add_run(taxon)

            photo = _first_photo(tree)
            if photo and not self._add_photo(document, photo, Cm(9), Cm(9)):
                document.add_paragraph("Fotografie není k dispozici")
            elif not photo:
                document.add_paragraph("Fotografie není k dispozici")

            document.add_heading("Komentář", level=2)
            comments = _comment_paragraphs(tree)
            if comments:
                for comment in comments:
                    document.add_paragraph(comment)
            else:
                document.add_paragraph("")

        document.save(output_path)

    def _add_photo(self, document, photo, max_width, max_height):
        try:
            from docx.shared import Pt

            with photo.photo.open("rb") as handle:
                image_bytes = handle.read()
            if not image_bytes:
                return False
            width, height = _fit_image_size(image_bytes, max_width, max_height)
            if width is None or height is None:
                return False
            document.add_picture(io.BytesIO(image_bytes), width=width, height=height)
            document.paragraphs[-1].paragraph_format.space_after = Pt(6)
            return True
        except Exception as exc:
            self.stderr.write(
                f"Skipping photo {photo.pk} for WorkRecord {photo.work_record_id}: {exc}"
            )
            return False
